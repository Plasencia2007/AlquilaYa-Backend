import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="888888"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def compile_markdown_to_docx(files_list, output_path):
    print("Iniciando compilación a Word (.docx)...")
    doc = Document()
    
    # Configuración de estilos por defecto (Times New Roman - IEEE)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Configurar márgenes de página
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)  # 1 inch
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)

    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    for file_path in files_list:
        if not os.path.exists(file_path):
            print(f"Advertencia: El archivo {file_path} no existe. Saltando...")
            continue
            
        print(f"Procesando {os.path.basename(file_path)}...")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
        for line in lines:
            stripped = line.strip()
            
            # Manejo de bloques de código
            if stripped.startswith("```"):
                if in_code_block:
                    # Guardar bloque de código acumulado
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Pt(18)
                    p.paragraph_format.right_indent = Pt(18)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    
                    # Añadir borde gris y fondo
                    pPr = p._p.get_or_add_pPr()
                    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="18" w:space="8" w:color="555555"/></w:pBdr>')
                    pPr.append(pBdr)
                    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
                    pPr.append(shd)
                    
                    code_text = "".join(code_content)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                    
                    code_content = []
                    in_code_block = False
                else:
                    in_code_block = True
                continue
                
            if in_code_block:
                code_content.append(line)
                continue

            # Manejo de tablas en Markdown
            if stripped.startswith("|"):
                # Ignorar la línea divisoria del header de markdown (ej: |:---|:---|)
                if re.match(r"^\|[\s|:-]*\|$", stripped):
                    continue
                # Recoger fila de la tabla
                cols = [c.strip() for c in stripped.split("|")[1:-1]]
                table_rows.append(cols)
                in_table = True
                continue
            else:
                if in_table and table_rows:
                    # Construir tabla acumulada en Word
                    rows_count = len(table_rows)
                    cols_count = len(table_rows[0])
                    
                    table = doc.add_table(rows=rows_count, cols=cols_count)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    set_table_borders(table)
                    
                    for r_idx, row_data in enumerate(table_rows):
                        row = table.rows[r_idx]
                        is_header = (r_idx == 0)
                        
                        # Impedir división de filas entre páginas
                        trPr = row._tr.get_or_add_trPr()
                        trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
                        
                        if is_header:
                            trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
                        
                        for c_idx, text in enumerate(row_data):
                            if c_idx < len(row.cells):
                                cell = row.cells[c_idx]
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                                
                                # Quitar formato de markdown de links o negritas dentro de las celdas
                                cell_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text) # limpia links
                                cell_text = cell_text.replace("**", "").replace("`", "")
                                
                                p = cell.paragraphs[0]
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                p.paragraph_format.space_before = Pt(2)
                                p.paragraph_format.space_after = Pt(2)
                                
                                run = p.add_run(cell_text)
                                run.font.name = 'Times New Roman'
                                if is_header:
                                    run.font.size = Pt(10)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                    set_cell_background(cell, "2B5C8F") # Azul corporativo
                                else:
                                    run.font.size = Pt(9.5)
                                    # Alternar color de fondo en filas
                                    if r_idx % 2 == 0:
                                        set_cell_background(cell, "F9FBFD")
                                    else:
                                        set_cell_background(cell, "FFFFFF")
                                        
                    # Espacio después de la tabla
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
                    table_rows = []
                    in_table = False

            # Ignorar líneas divisoras horizontales ---
            if stripped == "---":
                continue
                
            # Ignorar contenedores div de alineación de tablas
            if "<div" in stripped or "</div" in stripped:
                continue

            # Encabezados (#, ##, ###, ####)
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip('#'))
                title_text = stripped.lstrip('#').strip()
                # Limpiar negritas adicionales del título de markdown
                title_text = title_text.replace("**", "")
                
                p = doc.add_paragraph()
                p.paragraph_format.keep_with_next = True
                
                if level == 1:
                    p.paragraph_format.space_before = Pt(24)
                    p.paragraph_format.space_after = Pt(12)
                    run = p.add_run(title_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(16)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1E, 0x3F, 0x63)
                elif level == 2:
                    p.paragraph_format.space_before = Pt(18)
                    p.paragraph_format.space_after = Pt(8)
                    run = p.add_run(title_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(13)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x2B, 0x5C, 0x8F)
                elif level == 3:
                    p.paragraph_format.space_before = Pt(14)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(title_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                else:
                    p.paragraph_format.space_before = Pt(12)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run(title_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                continue

            # Listas desordenadas o con viñetas
            if stripped.startswith("* ") or stripped.startswith("- "):
                list_text = stripped[2:].strip()
                # Limpiar links y negritas
                list_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', list_text)
                list_text = list_text.replace("**", "").replace("`", "")
                
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Pt(24)
                
                run = p.add_run(list_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                continue

            # Párrafos regulares
            if stripped:
                # Limpiar sintaxis de links de markdown `[link](url)` -> `link`
                paragraph_text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', stripped)
                # Limpiar dobles asteriscos de negrita
                paragraph_text = paragraph_text.replace("**", "")
                # Limpiar comillas invertidas de código inline
                paragraph_text = paragraph_text.replace("`", "")
                
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                
                run = p.add_run(paragraph_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                
    # Guardar documento
    doc.save(output_path)
    print(f"Compilación terminada con éxito. Documento guardado en: {output_path}")

if __name__ == "__main__":
    docs = [
        "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/informe_seccion_6_c4.md",
        "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/informe_seccion_7_implementacion.md",
        "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/informe_seccion_8_seguridad.md",
        "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/informe_seccion_10_integracion.md",
        "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/informe_seccion_12_pruebas.md",
        "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/informe_seccion_13_conclusiones.md"
    ]
    out = "c:/Users/jhons/Documents/Proyectos/AlquilaYa-Backend/docs/AlquilaYa_Reporte_Completo.docx"
    compile_markdown_to_docx(docs, out)
